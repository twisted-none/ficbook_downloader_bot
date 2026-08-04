from __future__ import annotations

from datetime import datetime, timezone
from html import unescape
from io import BytesIO
from typing import Any

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer

from src.exporters.html_utils import html_to_text, html_to_text_blocks, normalize_inline_text, normalize_text, tag_text
from src.core.models import Story

ALLOWED_FORMATS: tuple[str, ...] = ("fb2", "epub", "txt", "docx", "pdf")
DEFAULT_FORMATS: tuple[str, ...] = ALLOWED_FORMATS
DOCX_PROPERTY_LIMIT = 255
PDF_FONT_PATHS = (
    "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
    "C:/Windows/Fonts/times.ttf",
    "C:/Windows/Fonts/timesbd.ttf",
)

CONTAINER_TAGS = {"article", "div", "section"}
INLINE_BOLD_TAGS = {"b", "strong"}
INLINE_ITALIC_TAGS = {"em", "i"}
SKIP_TAGS = {"audio", "iframe", "img", "script", "style", "video"}


def normalize_formats(values: list[str] | tuple[str, ...]) -> tuple[str, ...]:
    selected = set(values)
    formats = tuple(item for item in ALLOWED_FORMATS if item in selected)
    return formats or DEFAULT_FORMATS


def build_txt(story: Story) -> bytes:
    parts = [story.title, f"Автор: {story.author}", f"Ссылка: {story.source_url}"]
    if story.page_count:
        parts.append(f"Страниц: {story.page_count}")
    if story.word_count:
        parts.append(f"Слов: {story.word_count}")
    if story.pairings:
        parts.append(f"Пейринг: {_pairing_text(story)}")
    annotation = _html_to_text(story.annotation_html or story.description)
    if annotation:
        parts.extend(["Аннотация:", annotation])
    for index, chapter in enumerate(story.chapters, 1):
        title = chapter.title or f"Глава {index}"
        body = _html_to_text(chapter.html)
        parts.extend([title, body])
    return ("\n\n".join(part for part in parts if part.strip()) + "\n").encode("utf-8")


def build_docx(story: Story) -> bytes:
    document = Document()
    _configure_docx(document, story)
    _add_docx_page_footer(document)
    _add_docx_cover(document, story)
    document.add_heading(story.title, level=0)
    document.add_paragraph(f"Автор: {story.author}")
    document.add_paragraph(f"Ссылка: {story.source_url}")
    if story.page_count:
        document.add_paragraph(f"Страниц: {story.page_count}")
    if story.word_count:
        document.add_paragraph(f"Слов: {story.word_count}")
    if story.pairings:
        document.add_paragraph(f"Пейринг: {_pairing_text(story)}")
    _add_docx_toc(document, story)
    annotation = _html_to_blocks(story.annotation_html or story.description)
    if annotation:
        document.add_heading("Аннотация", level=1)
        _add_html_to_document(document, story.annotation_html or story.description)
    for index, chapter in enumerate(story.chapters, 1):
        document.add_heading(chapter.title or f"Глава {index}", level=1)
        _add_html_to_document(document, chapter.html)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf(story: Story) -> bytes:
    buffer = BytesIO()
    regular_font, bold_font = _register_pdf_fonts()
    styles = _pdf_styles(regular_font, bold_font)
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    flow: list[Any] = []
    _add_pdf_cover(flow, story)
    flow.extend(
        [
            Paragraph(_pdf_escape(story.title), styles["Title"]),
            Paragraph(f"Автор: {_pdf_escape(story.author)}", styles["Meta"]),
            Paragraph(f"Ссылка: {_pdf_escape(story.source_url)}", styles["Meta"]),
        ]
    )
    if story.page_count:
        flow.append(Paragraph(f"Страниц Ficbook: {_pdf_escape(story.page_count)}", styles["Meta"]))
    if story.word_count:
        flow.append(Paragraph(f"Слов: {_pdf_escape(story.word_count)}", styles["Meta"]))
    if story.pairings:
        flow.append(Paragraph(f"Пейринг: {_pdf_escape(_pairing_text(story))}", styles["Meta"]))
    flow.append(Spacer(1, 0.5 * cm))
    annotation = _html_to_blocks(story.annotation_html or story.description)
    if annotation:
        flow.append(Paragraph("Аннотация", styles["Heading1"]))
        _add_pdf_blocks(flow, annotation, styles)
    for index, chapter in enumerate(story.chapters, 1):
        if flow:
            flow.append(PageBreak())
        flow.append(Paragraph(_pdf_escape(chapter.title or f"Глава {index}"), styles["Heading1"]))
        _add_pdf_blocks(flow, _html_to_blocks(chapter.html), styles)
    document.build(flow, onFirstPage=_draw_pdf_page_number, onLaterPages=_draw_pdf_page_number)
    return buffer.getvalue()


def _configure_docx(document: Document, story: Story) -> None:
    core = document.core_properties
    core.title = _docx_property(story.title)
    core.author = _docx_property(story.author)
    core.subject = _docx_property(story.source_url)
    core.keywords = _docx_property(", ".join([*story.fandoms, *story.genres, *story.pairings]))
    pages = f"; Ficbook pages: {story.page_count}" if story.page_count else ""
    core.comments = _docx_property(f"Generated from {story.source_url}{pages}")
    core.created = datetime.now(timezone.utc)
    core.modified = datetime.now(timezone.utc)
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for style_name, size in {"Title": 20, "Heading 1": 16, "Heading 2": 14}.items():
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
    _ensure_paragraph_style(document, "Ficbook Quote", italic=True, left_indent=0.35)
    _ensure_paragraph_style(document, "Ficbook Note Heading", bold=True, space_before=6)


def _add_docx_cover(document: Document, story: Story) -> None:
    if not story.cover:
        return
    try:
        document.add_picture(BytesIO(story.cover.data), width=Inches(5.5))
        document.add_page_break()
    except (UnrecognizedImageError, ValueError):
        return


def _add_docx_page_footer(document: Document) -> None:
    for section in document.sections:
        paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = ""
        paragraph.add_run("Страница ")
        _add_docx_field(paragraph, "PAGE")
        paragraph.add_run(" из ")
        _add_docx_field(paragraph, "NUMPAGES")


def _add_docx_field(paragraph: Any, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)  # noqa: SLF001


def _register_pdf_fonts() -> tuple[str, str]:
    regular = PDF_FONT_PATHS[0]
    bold = PDF_FONT_PATHS[1]
    if _font_file_exists(regular) and _font_file_exists(bold):
        pdfmetrics.registerFont(TTFont("FicbookSerif", regular))
        pdfmetrics.registerFont(TTFont("FicbookSerif-Bold", bold))
        return "FicbookSerif", "FicbookSerif-Bold"
    if _font_file_exists(PDF_FONT_PATHS[2]) and _font_file_exists(PDF_FONT_PATHS[3]):
        pdfmetrics.registerFont(TTFont("FicbookSerif", PDF_FONT_PATHS[2]))
        pdfmetrics.registerFont(TTFont("FicbookSerif-Bold", PDF_FONT_PATHS[3]))
        return "FicbookSerif", "FicbookSerif-Bold"
    return "Times-Roman", "Times-Bold"


def _font_file_exists(path: str) -> bool:
    try:
        with open(path, "rb"):
            return True
    except OSError:
        return False


def _pdf_styles(regular_font: str, bold_font: str) -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "Title": ParagraphStyle(
            "FicbookTitle",
            parent=base["Title"],
            fontName=bold_font,
            fontSize=20,
            leading=24,
            spaceAfter=12,
        ),
        "Heading1": ParagraphStyle(
            "FicbookHeading1",
            parent=base["Heading1"],
            fontName=bold_font,
            fontSize=16,
            leading=20,
            spaceBefore=8,
            spaceAfter=8,
        ),
        "Body": ParagraphStyle(
            "FicbookBody",
            parent=base["BodyText"],
            fontName=regular_font,
            fontSize=11,
            leading=15,
            firstLineIndent=0.6 * cm,
            spaceAfter=6,
        ),
        "Meta": ParagraphStyle(
            "FicbookMeta",
            parent=base["BodyText"],
            fontName=regular_font,
            fontSize=10,
            leading=13,
            spaceAfter=4,
        ),
        "Footer": ParagraphStyle(
            "FicbookFooter",
            parent=base["BodyText"],
            alignment=TA_CENTER,
            fontName=regular_font,
            fontSize=9,
        ),
    }


def _add_pdf_cover(flow: list[Any], story: Story) -> None:
    if not story.cover:
        return
    try:
        image = Image(BytesIO(story.cover.data))
        max_width = 12 * cm
        max_height = 18 * cm
        scale = min(max_width / image.imageWidth, max_height / image.imageHeight, 1)
        image.drawWidth = image.imageWidth * scale
        image.drawHeight = image.imageHeight * scale
        image.hAlign = "CENTER"
        flow.extend([image, PageBreak()])
    except Exception:
        return


def _add_pdf_blocks(flow: list[Any], blocks: list[str], styles: dict[str, ParagraphStyle]) -> None:
    for block in blocks:
        flow.append(Paragraph(_pdf_escape(block), styles["Body"]))


def _draw_pdf_page_number(canvas: Any, document: Any) -> None:
    canvas.saveState()
    canvas.setFont("FicbookSerif" if "FicbookSerif" in pdfmetrics.getRegisteredFontNames() else "Times-Roman", 9)
    canvas.drawCentredString(A4[0] / 2, 1.15 * cm, f"Страница {document.page}")
    canvas.restoreState()


def _pdf_escape(text: str) -> str:
    return (
        unescape(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def _ensure_paragraph_style(
    document: Document,
    name: str,
    *,
    bold: bool = False,
    italic: bool = False,
    left_indent: float = 0.0,
    space_before: int = 0,
) -> None:
    try:
        style = document.styles[name]
    except KeyError:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.bold = bold
    style.font.italic = italic
    style.paragraph_format.left_indent = Inches(left_indent)
    style.paragraph_format.space_before = Pt(space_before)


def _docx_property(value: str) -> str:
    text = value.strip()
    if len(text) <= DOCX_PROPERTY_LIMIT:
        return text
    return text[: DOCX_PROPERTY_LIMIT - 1].rstrip() + "…"


def _pairing_text(story: Story) -> str:
    return ", ".join(story.pairings)


def _add_docx_toc(document: Document, story: Story) -> None:
    document.add_heading("Оглавление", level=1)
    if story.annotation_html or story.description:
        document.add_paragraph("Аннотация", style="List Bullet")
    for index, chapter in enumerate(story.chapters, 1):
        title = chapter.title or f"Глава {index}"
        document.add_paragraph(title, style="List Number")


def _html_to_text(html: str) -> str:
    return html_to_text(html)


def _html_to_blocks(html: str) -> list[str]:
    return html_to_text_blocks(html)


def _add_html_to_document(document: Document, html: str) -> None:
    if not html.strip():
        return
    soup = BeautifulSoup(html, "lxml")
    _add_blocks(document, soup.body or soup)


def _add_blocks(document: Document, container: Tag, *, quote: bool = False) -> None:
    for child in container.children:
        if isinstance(child, NavigableString):
            text = _normalize_text(str(child))
            if text:
                _add_paragraph(document, text, quote=quote)
            continue
        if child.name in SKIP_TAGS:
            continue
        if child.name in CONTAINER_TAGS:
            _add_blocks(document, child, quote=quote)
        elif child.name == "blockquote":
            _add_blocks(document, child, quote=True)
        elif child.name in {"ul", "ol"}:
            _add_list(document, child, ordered=child.name == "ol", quote=quote)
        elif child.name == "p":
            paragraph = _add_paragraph(document, "", quote=quote, note_heading=_is_note_heading(_tag_text(child)))
            _add_inline_runs(paragraph, child)
        elif child.name == "br":
            document.add_paragraph()
        else:
            paragraph = _add_paragraph(document, "", quote=quote)
            _add_inline_runs(paragraph, child)


def _add_list(document: Document, list_node: Tag, *, ordered: bool, quote: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for child in list_node.children:
        if not isinstance(child, Tag) or child.name != "li":
            continue
        paragraph = document.add_paragraph(style=style)
        if quote:
            paragraph.paragraph_format.left_indent = Inches(0.35)
        _add_inline_runs(paragraph, child, skip_nested_lists=True)
        for nested in child.find_all(["ul", "ol"], recursive=False):
            _add_list(document, nested, ordered=nested.name == "ol", quote=quote)


def _add_paragraph(document: Document, text: str, *, quote: bool = False, note_heading: bool = False):
    style = "Ficbook Note Heading" if note_heading else "Ficbook Quote" if quote else None
    paragraph = document.add_paragraph(text, style=style)
    if quote:
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.first_line_indent = Inches(0)
    return paragraph


def _add_inline_runs(
    paragraph: Any,
    node: Tag,
    *,
    bold: bool = False,
    italic: bool = False,
    skip_nested_lists: bool = False,
) -> None:
    for child in node.children:
        if isinstance(child, NavigableString):
            _add_run(paragraph, str(child), bold=bold, italic=italic)
            continue
        if child.name in SKIP_TAGS:
            continue
        if skip_nested_lists and child.name in {"ul", "ol"}:
            continue
        if child.name == "br":
            paragraph.add_run().add_break()
            continue
        _add_inline_runs(
            paragraph,
            child,
            bold=bold or child.name in INLINE_BOLD_TAGS,
            italic=italic or child.name in INLINE_ITALIC_TAGS,
            skip_nested_lists=skip_nested_lists,
        )


def _add_run(paragraph: Any, text: str, *, bold: bool, italic: bool) -> None:
    normalized = _normalize_inline_text(text)
    if not normalized:
        return
    run = paragraph.add_run(normalized)
    run.bold = bold
    run.italic = italic


def _tag_text(node: Tag) -> str:
    return tag_text(node)


def _is_note_heading(text: str) -> bool:
    return text.rstrip(":").lower() in {"примечания", "примечание автора", "примечания автора"}


def _normalize_text(text: str) -> str:
    return normalize_text(text)


def _normalize_inline_text(text: str) -> str:
    return normalize_inline_text(text)
